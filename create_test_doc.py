import docx
from docx.shared import Pt

# 1. 创建一个新的Word文档
doc = docx.Document()

# 2. 添加标题和说明文字
doc.add_heading('智能路径修正功能测试', level=1)
doc.add_paragraph('这是一个用于测试Word to Jupyter转换器智能路径修正功能的示例文档。')
doc.add_paragraph('下面的代码块包含一个相对路径，转换器应该能自动找到并修正它。')

# 3. 添加将被识别为“代码”的段落
# 我们将代码分成几行，模拟真实的书写习惯
code_lines = [
    "import pandas as pd",
    "import numpy as np",
    "# 程序会自动查找并修正下面的文件路径",
    "df = pd.read_csv('sample_data.csv')",
    "print('成功读取数据：')",
    "print(df.head())"
]

for line in code_lines:
    p = doc.add_paragraph(line)
    # 模拟代码样式
    font = p.runs[0].font
    font.name = 'Courier New'
    font.size = Pt(10)

# 4. 添加一些结尾的markdown文本
doc.add_paragraph('如果转换成功，生成的Jupyter Notebook应该能够直接运行并打印出CSV文件的头部数据。')

# 5. 保存文档
doc.save('test_path_correction.docx')

print("成功创建测试Word文档：'test_path_correction.docx'")